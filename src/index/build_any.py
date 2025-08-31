# src/index/build_any.py
from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional, Tuple

import typer
from rich import print
from leann import LeannBuilder

import requests
from readability import Document
from bs4 import BeautifulSoup
import fitz  # PyMuPDF

app = typer.Typer(add_completion=False)

UA = "leann-rag/0.1 (+local)"
_pdf_re = re.compile(r"\.pdf($|\?)", re.I)
_arxiv_id_re = re.compile(r"arxiv\.org/(?:pdf|abs|html)/([0-9]{4}\.[0-9]+)(v[0-9]+)?", re.I)

# ---------------- helpers: paths & names ----------------
def normalize_prefix(path_like: str) -> Path:
    """
    Accept .leann / .index / bare -> return base prefix Path (no suffix)
    """
    p = Path(path_like)
    p.parent.mkdir(parents=True, exist_ok=True)
    return p.with_suffix("") if p.suffix in {".leann", ".index"} else p

def slugify(s: str, max_len: int = 120) -> str:
    s = s.strip().lower()
    # replace non-word with hyphens
    s = re.sub(r"[^\w]+", "-", s)
    s = re.sub(r"-{2,}", "-", s).strip("-")
    if not s:
        s = "index"
    return s[:max_len]

def detect_arxiv_suffix(url_or_title: str) -> str:
    m = _arxiv_id_re.search(url_or_title)
    if not m:
        return ""
    base = m.group(1).replace(".", "-")
    ver = m.group(2) or ""
    return f"-{base}{ver}"

def unique_prefix(base_prefix: Path) -> Path:
    """
    Ensure we don't clobber an existing index. If <base>.index exists,
    append -1, -2, ... until unique.
    """
    if not base_prefix.with_suffix(".index").exists():
        return base_prefix
    i = 1
    while True:
        cand = base_prefix.with_name(f"{base_prefix.name}-{i}")
        if not cand.with_suffix(".index").exists():
            return cand
        i += 1

# ---------------- helpers: chunking ----------------
def chunk_text(text: str, size: int = 700, overlap: int = 100):
    if size <= 0:
        yield text
        return
    i, n = 0, len(text)
    step = max(1, size - overlap)
    while i < n:
        yield text[i : i + size]
        i += step

# ---------------- helpers: HTML & PDF ingestion ----------------
def fetch_url_text(url: str, timeout: int = 45) -> tuple[str, str]:
    r = requests.get(url, timeout=timeout, headers={"User-Agent": UA})
    r.raise_for_status()
    html = r.text
    doc = Document(html)
    title = (doc.short_title() or "").strip()
    content_html = doc.summary(html_partial=True)
    soup = BeautifulSoup(content_html, "html.parser")
    for bad in soup(["script", "style", "noscript", "header", "footer", "nav", "aside"]):
        bad.extract()
    text = soup.get_text(separator="\n")
    text = re.sub(r"\u00A0", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n\n", text).strip()
    return title or "", text

def fetch_pdf_text(url_or_path: str, timeout: int = 60) -> tuple[str, str]:
    if re.match(r"^https?://", url_or_path, re.I):
        r = requests.get(url_or_path, timeout=timeout, headers={"User-Agent": UA})
        r.raise_for_status()
        data = r.content
    else:
        data = Path(url_or_path).read_bytes()
    parts = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    text = "\n".join(parts)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    base = url_or_path.rstrip("/").split("/")[-1]
    title = base.replace(".pdf", "").replace("_", " ")
    return title, text.strip()

def is_pdf_url(u: str) -> bool:
    return bool(_pdf_re.search(u))

# ---------------- helpers: local files ----------------
def iter_files(root: Path, exts={".pdf", ".txt", ".md", ".docx"}):
    root = Path(root)
    if root.is_file():
        if root.suffix.lower() in exts:
            yield root
        return
    for p in root.rglob("*"):
        if p.is_file() and p.suffix.lower() in exts:
            yield p

def load_local_file(p: Path) -> dict:
    ext = p.suffix.lower()
    if ext in {".txt", ".md"}:
        return {"title": p.stem, "text": p.read_text(encoding="utf-8", errors="ignore")}
    if ext == ".pdf":
        title, text = fetch_pdf_text(str(p))
        return {"title": title or p.stem, "text": text}
    if ext == ".docx":
        import docx
        doc = docx.Document(str(p))
        text = "\n".join(para.text for para in doc.paragraphs)
        return {"title": p.stem, "text": text}
    raise ValueError(f"Unsupported file type: {p}")

# ---------------- CLI ----------------
@app.command()
def main(
    index_path: str = typer.Option(..., "--index-path", "-o", help="Output index path (.leann/.index/or bare prefix)"),
    url: Optional[List[str]] = typer.Option(None, "--url", help="One or more URLs (HTML or direct .pdf)"),
    file: Optional[List[str]] = typer.Option(None, "--file", help="One or more local files (.pdf/.txt/.md/.docx)"),
    dir_: Optional[List[str]] = typer.Option(None, "--dir", help="One or more directories to crawl"),
    chunk_size: int = typer.Option(700, help="Chunk size"),
    chunk_overlap: int = typer.Option(100, help="Chunk overlap"),
    auto_name: bool = typer.Option(True, "--auto-name/--no-auto-name", help="Derive output name from single-source title"),
):
    """
    Build a LEANN index from URLs, local files, or directories (graph-only, recompute embeddings at query time).
    With --auto-name (default), if there's exactly ONE input document, the index is named from its title
    (and arXiv id when detectable).
    """
    # 1) normalize initial prefix
    base_prefix = normalize_prefix(index_path)
    out_dir = base_prefix.parent

    # Gather inputs
    urls = url or []
    files = file or []
    dirs  = dir_ or []

    # Resolve and ingest sources while also capturing the FIRST doc's title+url for auto-naming
    builder = LeannBuilder(backend_name="hnsw", recompute=True)
    total_chunks = 0
    first_title: Optional[str] = None
    first_source_str: Optional[str] = None
    single_doc_count = 0  # count of logical docs (not chunks)

    # --- URLs ---
    for u in urls:
        print(f"[bold]Fetching[/]: {u}")
        try:
            title, text = fetch_pdf_text(u) if is_pdf_url(u) else fetch_url_text(u)
        except Exception as e:
            print(f"[red]Failed: {e}[/]")
            continue
        if not text:
            print(f"[yellow]No text extracted from {u}[/]")
            continue
        if first_title is None:
            first_title = title or ""
            first_source_str = u
        single_doc_count += 1
        for piece in chunk_text(text, size=chunk_size, overlap=chunk_overlap):
            builder.add_text(piece, metadata={"source": "url", "url": u, "title": title})
            total_chunks += 1

    # --- Files ---
    for f in files:
        rec = load_local_file(Path(f))
        title, text = rec["title"], rec["text"]
        if not text:
            print(f"[yellow]No text extracted from {f}[/]")
            continue
        if first_title is None:
            first_title = title or ""
            first_source_str = str(f)
        single_doc_count += 1
        for piece in chunk_text(text, size=chunk_size, overlap=chunk_overlap):
            builder.add_text(piece, metadata={"source": "file", "path": str(f), "title": title})
            total_chunks += 1

    # --- Directories ---
    for d in dirs:
        for p in iter_files(Path(d)):
            rec = load_local_file(p)
            title, text = rec["title"], rec["text"]
            if not text:
                print(f"[yellow]No text extracted from {p}[/]")
                continue
            if first_title is None:
                first_title = title or ""
                first_source_str = str(p)
            single_doc_count += 1
            for piece in chunk_text(text, size=chunk_size, overlap=chunk_overlap):
                builder.add_text(piece, metadata={"source": "file", "path": str(p), "title": title})
                total_chunks += 1

    if total_chunks == 0:
        print("[red]No content ingested[/]")
        raise typer.Exit(code=1)

    # 2) Autoname if exactly one logical document and flag is on
    final_prefix = base_prefix
    if auto_name and single_doc_count == 1 and first_title is not None:
        # derive from title + arXiv id if detectable
        slug = slugify(first_title)
        suffix = detect_arxiv_suffix(first_source_str or "")
        # only append the arXiv suffix if it isn't already present in the slug
        if suffix and not slug.endswith(suffix):
            slug = f"{slug}{suffix}"
        final_prefix = unique_prefix(out_dir / slug)


    # 3) Build index
    print(f"[bold]Building index[/] → {final_prefix}")
    builder.build_index(str(final_prefix))

    # 4) Report artifacts
    print(f"[green]Done[/]. Indexed {total_chunks} chunks.")
    print("Artifacts:")
    for sfx in [".index", ".meta.json", ".passages.idx", ".passages.jsonl"]:
        print(f" - {final_prefix.with_suffix(sfx)}")

if __name__ == "__main__":
    app()
