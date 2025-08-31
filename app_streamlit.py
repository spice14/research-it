# app_streamlit.py
from __future__ import annotations

import glob
import os
import re
import shutil
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import streamlit as st
from leann import LeannChat, LeannBuilder

# ---------------- Defaults ----------------
APP_ROOT = Path(__file__).resolve().parent
DEFAULT_INDEX = str(APP_ROOT / "indexes" / "arxiv-2307-09218.index")
DEFAULT_MODEL = "llama3.2:3b"
DEFAULT_TOP_K = 6
DEFAULT_NUM_CTX = 2048
SYSTEM = (
    "You are a helpful assistant. Use the retrieved CONTEXT faithfully. "
    "If the answer is not in the context, say you don't know."
)

UA = "leann-rag/0.1 (+local)"
_ARXIV_ID_RE = re.compile(r"arxiv\.org/(?:pdf|abs|html)/([0-9]{4}\.[0-9]+)(v[0-9]+)?", re.I)

# ---------------- Helpers (paths & discovery) ----------------
def resolve(path: str) -> Path:
    try:
        p = Path(path).expanduser()
        return p if p.is_absolute() else (Path.cwd() / p).resolve()
    except Exception:
        return Path(path)

def list_indexes(root: Path) -> List[str]:
    return sorted(glob.glob(str(root / "**/*.index"), recursive=True))

def expected_sidecars_for(index_path: Path) -> Tuple[Path, Path, Path]:
    base = index_path.name  # e.g., "foo.index"
    parent = index_path.parent
    meta     = parent / f"{base}.meta.json"
    pass_idx = parent / f"{base}.passages.idx"
    pass_jsl = parent / f"{base}.passages.jsonl"
    return meta, pass_idx, pass_jsl

def candidate_sources_for(index_path: Path) -> Dict[str, List[Path]]:
    parent = index_path.parent
    stem = index_path.stem  # "foo" for "foo.index"
    return {
        "meta": [
            parent / f"{stem}.meta.json",
            parent / f"{stem}.leann.meta.json",
        ],
        "pass_idx": [
            parent / f"{stem}.passages.idx",
            parent / f"{stem}.leann.passages.idx",
        ],
        "pass_jsl": [
            parent / f"{stem}.passages.jsonl",
            parent / f"{stem}.leann.passages.jsonl",
        ],
    }

def ensure_sidecars(index_path: Path) -> Dict[str, bool]:
    expected_meta, expected_idx, expected_jsl = expected_sidecars_for(index_path)
    status = {
        "expected_meta_exists": expected_meta.exists(),
        "expected_idx_exists": expected_idx.exists(),
        "expected_jsl_exists": expected_jsl.exists(),
    }
    if all(status.values()):
        return status

    cands = candidate_sources_for(index_path)

    def create_if_missing(expected: Path, sources: List[Path]):
        if expected.exists():
            return
        for src in sources:
            if src.exists():
                shutil.copy2(src, expected)
                break

    create_if_missing(expected_meta, cands["meta"])
    create_if_missing(expected_idx,  cands["pass_idx"])
    create_if_missing(expected_jsl,  cands["pass_jsl"])

    status["expected_meta_exists"] = expected_meta.exists()
    status["expected_idx_exists"]  = expected_idx.exists()
    status["expected_jsl_exists"]  = expected_jsl.exists()
    return status

def normalize_index_prefix(user_path: str) -> Path:
    """Accept .leann/.index/bare; return a base prefix Path (no suffix)."""
    p = Path(user_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    if p.suffix in {".leann", ".index"}:
        return p.with_suffix("")
    return p

def slugify(s: str, max_len: int = 120) -> str:
    s = s.strip().lower()
    s = re.sub(r"[^\w]+", "-", s)
    s = re.sub(r"-{2,}", "-", s).strip("-")
    if not s:
        s = "index"
    return s[:max_len]

def detect_arxiv_suffix(url_or_title: str) -> str:
    m = _ARXIV_ID_RE.search(url_or_title)
    if not m:
        return ""
    base = m.group(1).replace(".", "-")
    ver = m.group(2) or ""
    return f"-{base}{ver}"

def unique_prefix(base_prefix: Path) -> Path:
    """If <base>.index exists, append -1, -2, ... until unique."""
    if not base_prefix.with_suffix(".index").exists():
        return base_prefix
    i = 1
    while True:
        cand = base_prefix.with_name(f"{base_prefix.name}-{i}")
        if not cand.with_suffix(".index").exists():
            return cand
        i += 1

# ---------------- Minimal ingestion utils (URL/HTML, PDF, files, chunk) ----------------
def fetch_url_text(url: str) -> tuple[str, str]:
    import requests
    from readability import Document
    from bs4 import BeautifulSoup

    r = requests.get(url, timeout=45, headers={"User-Agent": UA})
    r.raise_for_status()
    html = r.text
    doc = Document(html)
    title = (doc.short_title() or "").strip()
    content_html = doc.summary(html_partial=True)
    soup = BeautifulSoup(content_html, "html.parser")
    for bad in soup(["script", "style", "noscript", "header", "footer", "nav", "aside"]):
        bad.extract()
    text = soup.get_text(separator="\n")
    text = re.sub(r"\u00A0", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n\n", text).strip()
    return title, text

def fetch_pdf_text(url_or_path: str) -> tuple[str, str]:
    import requests, fitz
    if re.match(r"^https?://", url_or_path, re.I):
        r = requests.get(url_or_path, timeout=60, headers={"User-Agent": UA})
        r.raise_for_status()
        data = r.content
    else:
        data = Path(url_or_path).read_bytes()
    parts = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    text = "\n".join(parts)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    base = url_or_path.rstrip("/").split("/")[-1]
    title = base.replace(".pdf", "").replace("_", " ")
    return title or "document", text.strip()

def is_pdf_url(u: str) -> bool:
    return bool(re.search(r"\.pdf($|\?)", u, re.I))

def iter_files(root: Path, exts={".pdf", ".txt", ".md", ".docx"}):
    root = Path(root)
    if root.is_file():
        if root.suffix.lower() in exts:
            yield root
        return
    for p in root.rglob("*"):
        if p.is_file() and p.suffix.lower() in exts:
            yield p

def load_local_file(p: Path) -> Dict[str, str]:
    ext = p.suffix.lower()
    if ext in {".txt", ".md"}:
        return {"title": p.stem, "text": p.read_text(encoding="utf-8", errors="ignore")}
    if ext == ".pdf":
        title, text = fetch_pdf_text(str(p))
        return {"title": title or p.stem, "text": text}
    if ext == ".docx":
        import docx
        doc = docx.Document(str(p))
        text = "\n".join(para.text for para in doc.paragraphs)
        return {"title": p.stem, "text": text}
    raise ValueError(f"Unsupported file type: {p}")

def chunk_text(text: str, size: int = 600, overlap: int = 80):
    if size <= 0:
        yield text
        return
    n = len(text); i = 0
    step = max(1, size - overlap)
    while i < n:
        yield text[i:i+size]
        i += step

# ---------------- Builders used by UI ----------------
def build_index_from_inputs(
    urls: List[str],
    files: List[Path],
    dirs: List[Path],
    index_path_input: str,
    chunk_size: int,
    chunk_overlap: int,
    auto_name: bool = True,
) -> Path:
    """
    Build a LEANN index (graph-only, recompute=True) from mixed sources.
    Returns absolute path to the .index file.
    """
    index_prefix = normalize_index_prefix(index_path_input)
    index_prefix.parent.mkdir(parents=True, exist_ok=True)

    data_dir = APP_ROOT / "data"
    data_dir.mkdir(parents=True, exist_ok=True)

    builder = LeannBuilder(backend_name="hnsw", recompute=True)
    total_chunks = 0
    first_title: Optional[str] = None
    first_source_str: Optional[str] = None
    logical_docs = 0

    # URLs
    for u in urls:
        if not u.strip():
            continue
        if is_pdf_url(u):
            title, text = fetch_pdf_text(u)
        else:
            title, text = fetch_url_text(u)
        if not text:
            continue
        # save cleaned text (optional)
        safe = (title or "page").replace("/", "_")[:120] or "doc"
        (data_dir / f"{safe}.txt").write_text(text, encoding="utf-8")
        if first_title is None:
            first_title, first_source_str = title or "", u
        logical_docs += 1
        for piece in chunk_text(text, size=chunk_size, overlap=chunk_overlap):
            builder.add_text(piece, metadata={"source": "url", "url": u, "title": title})
            total_chunks += 1

    # Files
    for f in files:
        rec = load_local_file(f)
        title, text = rec["title"], rec["text"]
        if not text:
            continue
        (data_dir / f"{f.stem}.txt").write_text(text, encoding="utf-8")
        if first_title is None:
            first_title, first_source_str = title or "", str(f)
        logical_docs += 1
        for piece in chunk_text(text, size=chunk_size, overlap=chunk_overlap):
            builder.add_text(piece, metadata={"source": "file", "path": str(f), "title": title})
            total_chunks += 1

    # Dirs
    for d in dirs:
        for p in iter_files(d):
            rec = load_local_file(p)
            title, text = rec["title"], rec["text"]
            if not text:
                continue
            (data_dir / f"{p.stem}.txt").write_text(text, encoding="utf-8")
            if first_title is None:
                first_title, first_source_str = title or "", str(p)
            logical_docs += 1
            for piece in chunk_text(text, size=chunk_size, overlap=chunk_overlap):
                builder.add_text(piece, metadata={"source": "file", "path": str(p), "title": title})
                total_chunks += 1

    if total_chunks == 0:
        raise RuntimeError("No content ingested; nothing to index.")

    # Optional auto-name when exactly one logical doc
    final_prefix = index_prefix
    if auto_name and logical_docs == 1 and first_title is not None:
        slug = slugify(first_title)
        suffix = detect_arxiv_suffix(first_source_str or "")
        if suffix and not slug.endswith(suffix):
            slug = f"{slug}{suffix}"
        final_prefix = unique_prefix(index_prefix.parent / slug)

    # build index
    builder.build_index(str(final_prefix))

    index_file = final_prefix.with_suffix(".index")
    ensure_sidecars(index_file)
    return index_file.resolve()

# ---------------- Chat cache ----------------
@st.cache_resource(show_spinner=False)
def get_chat(resolved_index_path: str, model: str, num_ctx: int) -> LeannChat:
    ix = Path(resolved_index_path)
    if ix.suffix != ".index":
        raise ValueError(f"Expected a .index file, got: {ix}")
    status = ensure_sidecars(ix)
    if not all(status.values()):
        missing = [k for k, ok in status.items() if not ok]
        raise FileNotFoundError(
            f"Missing required LEANN sidecar(s) for {ix.name}: {', '.join(missing)}."
        )
    return LeannChat(
        resolved_index_path,
        llm_config={"type": "ollama", "model": model, "num_ctx": num_ctx},
        system_prompt=SYSTEM,
    )

def run_rag(chat: LeannChat, question: str, top_k: int) -> Dict[str, Any]:
    out = chat.ask(question, top_k=top_k)
    if isinstance(out, dict):
        return {"answer": out.get("text", ""), "sources": out.get("sources", [])}
    return {"answer": str(out), "sources": []}

def source_title(meta: Dict) -> str:
    return meta.get("filename") or meta.get("title") or meta.get("path") or meta.get("url") or "source"

def trim_text(t: str, limit: int = 800) -> str:
    if not isinstance(t, str):
        return ""
    return (t if len(t) <= limit else t[:limit] + "…").strip()

# ---------------- UI ----------------
st.set_page_config(page_title="LEANN RAG (Local)", page_icon="🧠", layout="wide")
st.title("🧠 LEANN RAG — Local (Ollama)")

with st.sidebar:
    st.header("Settings")

    # Existing index picker
    indexes_root = APP_ROOT / "indexes"
    found = list_indexes(indexes_root) if indexes_root.exists() else []
    picked = st.selectbox(
        "Pick an existing LEANN index (.index)",
        options=["(type path manually)"] + found,
        index=0,
    )
    raw_index_path = DEFAULT_INDEX if picked == "(type path manually)" else picked
    user_index_path = st.text_input("Index (.index) path", value=raw_index_path)

    model = st.text_input("Ollama model", value=DEFAULT_MODEL)
    top_k = st.slider("Top-K (retrieval)", 1, 20, value=DEFAULT_TOP_K, step=1)
    num_ctx = st.slider("LLM context (num_ctx)", 512, 8192, value=DEFAULT_NUM_CTX, step=256)

    st.divider()
    st.subheader("Build new index")

    tabs = st.tabs(["🌐 URL", "📄 Upload files", "📁 Folder"])

    # --- URL tab ---
    with tabs[0]:
        url_to_index = st.text_input("URL (HTML or direct PDF)", value="https://arxiv.org/pdf/2307.09218v3.pdf")
        out_path_url = st.text_input("Output index prefix (.leann/.index/or bare)", value=str(APP_ROOT / "indexes" / "new-url-index.leann"))
        cs_url = st.number_input("Chunk size", min_value=100, max_value=4000, value=700, step=50, key="cs_url")
        co_url = st.number_input("Chunk overlap", min_value=0, max_value=1000, value=100, step=10, key="co_url")
        auto_url = st.checkbox("Auto-name from title/arXiv id (single source)", value=True, key="auto_url")
        build_url_btn = st.button("Build from URL", type="primary", use_container_width=True)

    # --- Upload tab ---
    with tabs[1]:
        uploads = st.file_uploader(
            "Upload PDF / TXT / MD / DOCX (one or more)",
            type=["pdf", "txt", "md", "docx"],
            accept_multiple_files=True,
        )
        out_path_upload = st.text_input("Output index prefix (.leann/.index/or bare)", value=str(APP_ROOT / "indexes" / "uploaded-docs.leann"))
        cs_up = st.number_input("Chunk size ", min_value=100, max_value=4000, value=700, step=50, key="cs_up")
        co_up = st.number_input("Chunk overlap ", min_value=0, max_value=1000, value=100, step=10, key="co_up")
        auto_up = st.checkbox("Auto-name from title/arXiv id (single source)", value=True, key="auto_up")
        build_upload_btn = st.button("Build from uploads", type="primary", use_container_width=True)

    # --- Folder tab ---
    with tabs[2]:
        folder_path = st.text_input("Folder path to crawl (PDF/TXT/MD/DOCX)", value=str(APP_ROOT / "data"))
        out_path_dir = st.text_input("Output index prefix (.leann/.index/or bare)", value=str(APP_ROOT / "indexes" / "folder-index.leann"))
        cs_dir = st.number_input("Chunk size  ", min_value=100, max_value=4000, value=700, step=50, key="cs_dir")
        co_dir = st.number_input("Chunk overlap  ", min_value=0, max_value=1000, value=100, step=10, key="co_dir")
        auto_dir = st.checkbox("Auto-name from title/arXiv id (single source)", value=False, key="auto_dir")
        build_dir_btn = st.button("Build from folder", type="primary", use_container_width=True)

    cold_start = st.button("Re-init pipeline")

    # Diagnostics
    st.divider()
    st.caption("Diagnostics")
    resolved = resolve(user_index_path)
    exists = resolved.exists()
    size = resolved.stat().st_size if exists else 0
    st.write(f"**CWD:** `{os.getcwd()}`")
    st.write(f"**App root:** `{APP_ROOT}`")
    st.write("**Resolved index path:**")
    st.code(str(resolved))
    st.write(f"**Exists:** {'✅' if exists else '❌'}")
    st.write(f"**Size:** {size} bytes")

    if exists and resolved.suffix == ".index":
        em, ei, ej = expected_sidecars_for(resolved)
        st.write("**Expected sidecars:**")
        st.write(f"- {em.name}: {'✅' if em.exists() else '❌'}")
        st.write(f"- {ei.name}: {'✅' if ei.exists() else '❌'}")
        st.write(f"- {ej.name}: {'✅' if ej.exists() else '❌'}")

    if indexes_root.exists():
        st.write("**Discovered under ./indexes:**")
        if found:
            for p in found[:20]:
                st.code(p)
            if len(found) > 20:
                st.caption(f"... and {len(found)-20} more")
        else:
            st.caption("(no .index files found)")

# ---- Build actions (URL / Upload / Folder) ----
def after_success(new_ix: Path):
    global user_index_path, resolved
    st.success(f"Index built: {new_ix}")
    user_index_path = str(new_ix)
    resolved = new_ix
    if "messages" in st.session_state:
        st.session_state["messages"].clear()
    get_chat.clear()

if 'resolved' not in locals():
    resolved = resolve(user_index_path)

if build_url_btn:
    with st.spinner("Fetching & indexing URL…"):
        try:
            new_ix = build_index_from_inputs(
                urls=[url_to_index.strip()],
                files=[],
                dirs=[],
                index_path_input=out_path_url.strip(),
                chunk_size=int(cs_url),
                chunk_overlap=int(co_url),
                auto_name=bool(auto_url),
            )
            after_success(new_ix)
        except Exception as e:
            st.error("Failed to build index from URL:")
            st.exception(e)

if build_upload_btn:
    with st.spinner("Indexing uploads…"):
        try:
            upload_dir = APP_ROOT / "uploads"
            upload_dir.mkdir(parents=True, exist_ok=True)
            saved_files: List[Path] = []
            for uf in uploads or []:
                p = upload_dir / uf.name
                with open(p, "wb") as fp:
                    fp.write(uf.getbuffer())
                saved_files.append(p)
            if not saved_files:
                st.warning("No files uploaded.")
            else:
                new_ix = build_index_from_inputs(
                    urls=[],
                    files=saved_files,
                    dirs=[],
                    index_path_input=out_path_upload.strip(),
                    chunk_size=int(cs_up),
                    chunk_overlap=int(co_up),
                    auto_name=bool(auto_up),
                )
                after_success(new_ix)
        except Exception as e:
            st.error("Failed to build index from uploaded files:")
            st.exception(e)

if build_dir_btn:
    with st.spinner("Indexing folder…"):
        try:
            new_ix = build_index_from_inputs(
                urls=[],
                files=[],
                dirs=[Path(folder_path.strip())],
                index_path_input=out_path_dir.strip(),
                chunk_size=int(cs_dir),
                chunk_overlap=int(co_dir),
                auto_name=bool(auto_dir),
            )
            after_success(new_ix)
        except Exception as e:
            st.error("Failed to build index from folder:")
            st.exception(e)

# (Re)initialize pipeline
chat = None
init_error: Exception | None = None
if Path(user_index_path).exists():
    try:
        if cold_start:
            get_chat.clear()
        chat = get_chat(str(resolved), DEFAULT_MODEL if not model else model, num_ctx)
    except Exception as e:
        init_error = e
else:
    st.sidebar.warning("Index not found at the resolved path above. Use a .index file path or build one.")

# show init errors
if init_error:
    st.error("Failed to initialize LeannChat. See details below.")
    st.exception(init_error)

# keep chat history
if "messages" not in st.session_state:
    st.session_state["messages"] = []

# show history
for m in st.session_state["messages"]:
    with st.chat_message(m["role"]):
        st.markdown(m["content"])
        if m["role"] == "assistant" and m.get("sources"):
            with st.expander("Sources"):
                for i, s in enumerate(m["sources"], 1):
                    meta = s.get("metadata", {}) if isinstance(s, dict) else {}
                    title = source_title(meta)
                    st.markdown(f"**{i}. {title}**")
                    if "url" in meta:
                        st.markdown(f"[link]({meta['url']})")
                    if "path" in meta:
                        st.code(meta["path"])
                    snippet = s.get("text", "")
                    if snippet:
                        st.write(trim_text(snippet))
                    st.divider()

# input & inference
user_q = st.chat_input("Ask a question about your indexed documents…")
if user_q:
    st.session_state["messages"].append({"role": "user", "content": user_q})
    with st.chat_message("user"):
        st.markdown(user_q)

    if chat is None:
        with st.chat_message("assistant"):
            if init_error:
                st.error("Pipeline failed to initialize:")
                st.exception(init_error)
            else:
                st.error("Index not found / invalid. Pick a .index file or build one in the sidebar.")
    else:
        with st.chat_message("assistant"):
            with st.spinner("Thinking…"):
                try:
                    out = run_rag(chat, user_q, top_k=DEFAULT_TOP_K if top_k is None else top_k)
                    answer = out["answer"] or "_(No answer returned)_"
                    sources = out.get("sources", [])
                    st.markdown(answer)
                    if sources:
                        with st.expander("Sources"):
                            for i, s in enumerate(sources, 1):
                                meta = s.get("metadata", {}) if isinstance(s, dict) else {}
                                title = source_title(meta)
                                st.markdown(f"**{i}. {title}**")
                                if "url" in meta:
                                    st.markdown(f"[link]({meta['url']})")
                                if "path" in meta:
                                    st.code(meta["path"])
                                snippet = s.get("text", "")
                                if snippet:
                                    st.write(trim_text(snippet))
                                st.divider()
                    st.session_state["messages"].append(
                        {"role": "assistant", "content": answer, "sources": sources}
                    )
                except Exception as e:
                    st.error("Error during RAG:")
                    st.exception(e)
